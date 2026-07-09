import streamlit as st
import pdfplumber
import requests

from docx import Document

# ==================================================
# CONFIGURAÇÃO
# ==================================================

def pagina_penajusta():
    st.title("Clipping de Diários")
    st.caption("....")
    
    PALAVRAS = [
        "SEPPEN",
        "PENA JUSTA",
        "CEPPRJ",
        "GMF",
        "superlotação prisional"
    ]

# ==================================================
# FUNÇÕES
# ==================================================

    def baixar_pdf(url, nome_arquivo):
    
        resposta = requests.get(
            url,
            timeout=60
        )
    
        resposta.raise_for_status()
    
        with open(nome_arquivo, "wb") as f:
            f.write(resposta.content)
    
        return nome_arquivo
    
    
    def extrair_texto(pdf_path):
    
        texto = ""
    
        with pdfplumber.open(pdf_path) as pdf:
    
            for pagina in pdf.pages:
    
                conteudo = pagina.extract_text()
    
                if conteudo:
                    texto += conteudo + "\n"
    
        return texto
    
    
    def obter_linhas(texto):
    
        return [
            linha.strip()
            for linha in texto.split("\n")
            if linha.strip()
        ]
    
    
    def localizar_ocorrencias(linhas, palavras):
    
        resultados = []
    
        for indice, linha in enumerate(linhas):
    
            for palavra in palavras:
    
                if palavra.lower() in linha.lower():
    
                    resultados.append(
                        {
                            "linha": indice,
                            "palavra": palavra,
                            "texto": linha
                        }
                    )
    
        return resultados
    
    
    def gerar_word(blocos):
    
        doc = Document()
    
        doc.add_heading(
            "CLIPPING DIÁRIO",
            level=0
        )
    
        for i, bloco in enumerate(
            blocos,
            start=1
        ):
    
            doc.add_heading(
                f"Ocorrência {i}",
                level=1
            )
    
            doc.add_paragraph(
                f"Documento: {bloco['documento']}"
            )
    
            doc.add_paragraph(
                f"Palavra-chave: {bloco['palavra']}"
            )
    
            doc.add_paragraph(
                bloco["conteudo"]
            )
    
        caminho = "clipping.docx"
    
        doc.save(caminho)
    
        return caminho


# ==================================================
# SESSION STATE
# ==================================================

if "clipping" not in st.session_state:

    st.session_state.clipping = []

# ==================================================
# INTERFACE
# ==================================================

st.title(
    "📑 Clipping de Diários Oficiais"
)

col1, col2 = st.columns(2)

with col1:

    url_cnj = st.text_input(
        "CNJ"
    )

    url_cnmp = st.text_input(
        "CNMP"
    )

with col2:

    url_tjrj = st.text_input(
        "TJRJ"
    )

    url_ioerj = st.text_input(
        "IOERJ"
    )

processar = st.button(
    "Processar Diários"
)

# ==================================================
# PROCESSAMENTO
# ==================================================

if processar:

    documentos = [

        ("CNJ", url_cnj),
        ("CNMP", url_cnmp),
        ("TJRJ", url_tjrj),
        ("IOERJ", url_ioerj)

    ]

    for nome_doc, url in documentos:

        if not url.strip():
            continue

        try:

            caminho_pdf = f"{nome_doc}.pdf"

            baixar_pdf(
                url,
                caminho_pdf
            )

            texto = extrair_texto(
                caminho_pdf
            )

            linhas = obter_linhas(
                texto
            )

            ocorrencias = localizar_ocorrencias(
                linhas,
                PALAVRAS
            )

            st.header(nome_doc)

            st.success(
                f"{len(ocorrencias)} ocorrência(s) encontrada(s)"
            )

            if len(ocorrencias) == 0:

                st.warning(
                    "Nenhuma palavra localizada."
                )

            else:

                for indice, item in enumerate(
                    ocorrencias
                ):

                    with st.expander(
                        f"{item['palavra']} - Linha {item['linha']}"
                    ):

                        linha_chave = item["linha"]

                        inicio = max(
                            0,
                            linha_chave - 20
                        )

                        fim = min(
                            len(linhas),
                            linha_chave + 30
                        )

                        contexto = "\n".join(
                            linhas[inicio:fim]
                        )

                        st.text_area(
                            "Contexto",
                            contexto,
                            height=250,
                            key=f"{nome_doc}_{indice}"
                        )

                        if st.button(
                            f"Adicionar ao clipping",
                            key=f"add_{nome_doc}_{indice}"
                        ):

                            st.session_state.clipping.append(

                                {
                                    "documento": nome_doc,
                                    "palavra": item["palavra"],
                                    "conteudo": contexto
                                }

                            )

                            st.success(
                                "Adicionado."
                            )

        except Exception as erro:

            st.error(
                f"Erro em {nome_doc}: {erro}"
            )

# ==================================================
# CLIPPING
# ==================================================

st.divider()

st.subheader(
    "Clipping Atual"
)

if len(st.session_state.clipping) == 0:

    st.info(
        "Nenhum trecho selecionado."
    )

else:

    for i, bloco in enumerate(
        st.session_state.clipping,
        start=1
    ):

        with st.expander(
            f"Trecho {i}"
        ):

            st.write(
                f"Documento: {bloco['documento']}"
            )

            st.write(
                f"Palavra: {bloco['palavra']}"
            )

            st.text(
                bloco["conteudo"]
            )

# ==================================================
# WORD
# ==================================================

if len(st.session_state.clipping) > 0:

    if st.button(
        "Gerar Word"
    ):

        arquivo_word = gerar_word(
            st.session_state.clipping
        )

        with open(
            arquivo_word,
            "rb"
        ) as f:

            st.download_button(
                "📥 Baixar Clipping",
                data=f,
                file_name="clipping.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
